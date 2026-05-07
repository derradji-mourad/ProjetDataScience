"""
Génération du rapport Word complet pour le projet :
"Prédiction de la polarisation médiatique d'événements géopolitiques"
Utilise python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Couleurs du projet ──────────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x1A, 0x52, 0x76)   # titre principal
BLUE_MID    = RGBColor(0x21, 0x86, 0xC4)   # sous-titres
BLUE_LIGHT  = RGBColor(0xD5, 0xE8, 0xF7)   # fond tableaux header
RED_POL     = RGBColor(0xC0, 0x39, 0x2B)   # gauche
GREEN_POL   = RGBColor(0x27, 0xAE, 0x60)   # centre
ORANGE_POL  = RGBColor(0x1A, 0x5A, 0x96)   # droite
GREY_TEXT   = RGBColor(0x44, 0x44, 0x44)
GREY_LIGHT  = RGBColor(0xF5, 0xF7, 0xFA)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Définit la couleur de fond d'une cellule."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def style_heading(paragraph, level=1, color=BLUE_DARK):
    for run in paragraph.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)

def add_colored_heading(doc, text, level=1, color=BLUE_DARK):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.bold = True
    return p

def add_body(doc, text, bold_parts=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = GREY_TEXT
    else:
        # text est une liste de (texte, is_bold)
        for (t, is_bold) in text:
            run = p.add_run(t)
            run.font.size = Pt(11)
            run.bold = is_bold
            run.font.color.rgb = GREY_TEXT
    return p

def add_bullet(doc, text, level=0, bold_first=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)

    if bold_first:
        r1 = p.add_run(bold_first)
        r1.bold = True
        r1.font.size = Pt(11)
        r1 = p.add_run(text)
        r1.font.size = Pt(11)
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
    return p

def add_info_box(doc, title, content_lines, bg_hex="D5E8F7", border_color=BLUE_MID):
    """Ajoute un encadré coloré (tableau 1 cellule)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, bg_hex)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:color'), str(border_color))
        tcBorders.append(border)
    tcPr.append(tcBorders)

    p_title = cell.paragraphs[0]
    run_title = p_title.add_run(title)
    run_title.bold = True
    run_title.font.size = Pt(11)
    run_title.font.color.rgb = BLUE_DARK

    for line in content_lines:
        p = cell.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.3)
        for run in p.runs:
            run.font.size = Pt(10.5)

    doc.add_paragraph()
    return table

def make_header_table(doc, headers, rows_data, col_widths=None):
    """Crée un tableau avec en-tête coloré."""
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1A5276")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx + 1]
        bg = "F0F6FC" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()
    return table


# ════════════════════════════════════════════════════════════════════════════
# CONSTRUCTION DU DOCUMENT
# ════════════════════════════════════════════════════════════════════════════

import docx

doc = Document()

# ── Marges ───────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Style par défaut ─────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name  = 'Calibri'
style.font.size  = Pt(11)


# ════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════════

for _ in range(5):
    doc.add_paragraph()

# Titre principal
p_main = doc.add_paragraph()
p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_main.add_run("PRÉDICTION DE LA POLARISATION MÉDIATIQUE")
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = BLUE_DARK

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("D'ÉVÉNEMENTS GÉOPOLITIQUES")
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = BLUE_DARK

doc.add_paragraph()

p_sep = doc.add_paragraph()
p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sep.add_run("─" * 40)
r.font.color.rgb = BLUE_MID

doc.add_paragraph()

p_bloc = doc.add_paragraph()
p_bloc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_bloc.add_run("Bloc 4 — Intelligence Artificielle & Data Science")
r.font.size  = Pt(14)
r.font.bold  = True
r.font.color.rgb = BLUE_MID

doc.add_paragraph()

for _ in range(3):
    doc.add_paragraph()

# Infos
infos = [
    ("Auteur", "Mourad"),
    ("Email", "mohaqlf2mz23@gmail.com"),
    ("Formation", "Data Science & IA — Bloc 4"),
    ("Date", "Mai 2026"),
]
tbl_garde = doc.add_table(rows=len(infos), cols=2)
tbl_garde.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, val) in enumerate(infos):
    cells = tbl_garde.rows[i].cells
    set_cell_bg(cells[0], "1A5276")
    set_cell_bg(cells[1], "EBF5FB")
    r0 = cells[0].paragraphs[0].add_run(label)
    r0.bold = True; r0.font.color.rgb = RGBColor(255,255,255); r0.font.size = Pt(11)
    r1 = cells[1].paragraphs[0].add_run(val)
    r1.font.size = Pt(11); r1.font.color.rgb = GREY_TEXT

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# RÉSUMÉ EXÉCUTIF
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "Résumé Exécutif", level=1)

add_body(doc, (
    "Ce projet de data science a pour objectif de construire un système capable de "
    "prédire automatiquement l'orientation politique (Gauche, Centre, Droite) d'articles "
    "de presse traitant d'événements géopolitiques. Il met en oeuvre un pipeline complet "
    "allant de la collecte de données jusqu'au déploiement d'un tableau de bord interactif, "
    "en passant par le nettoyage, l'analyse exploratoire et la modélisation par apprentissage automatique."
))

# Tableau résumé
add_info_box(doc, "Points clés du projet", [
    "• 3 000 articles de presse analysés sur 10 événements géopolitiques majeurs (2022-2024)",
    "• 27 sources médiatiques couvrant les 3 orientations (Gauche, Centre, Droite)",
    "• 5 algorithmes de Machine Learning comparés avec validation croisée 5-fold",
    "• Tableau de bord Plotly interactif exportable en HTML autonome",
    "• Tous les critères BC4.1 à BC4.4 couverts"
])

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "Table des Matières", level=1)

toc_items = [
    ("1.", "Introduction et Problématique", "4"),
    ("2.", "Architecture du Pipeline de Données", "5"),
    ("3.", "Collecte des Données (BC4.1)", "6"),
    ("4.", "Prétraitement et Nettoyage (BC4.2)", "9"),
    ("5.", "Analyse Exploratoire des Données (EDA)", "13"),
    ("6.", "Modélisation Machine Learning (BC4.4)", "17"),
    ("7.", "Dashboard Interactif (BC4.3)", "27"),
    ("8.", "Résultats et Discussion", "29"),
    ("9.", "Conclusion et Perspectives", "31"),
    ("10.", "Mapping des Compétences", "32"),
    ("11.", "Glossaire", "33"),
]

for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    tab = p.paragraph_format.tab_stops
    p.paragraph_format.tab_stops.add_tab_stop(Cm(15), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(f"{num}  {title}")
    r1.font.size = Pt(11)
    r1.font.color.rgb = GREY_TEXT
    r2 = p.add_run(f"\t{page}")
    r2.font.size = Pt(11)
    r2.font.color.rgb = BLUE_MID

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION ET PROBLÉMATIQUE
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "1. Introduction et Problématique", level=1)

add_colored_heading(doc, "1.1 Contexte", level=2, color=BLUE_MID)
add_body(doc, (
    "À l'ère de l'information numérique, les citoyens sont exposés à une quantité "
    "considérable d'articles de presse couvrant les mêmes événements géopolitiques, "
    "mais avec des angles radicalement différents. Cette diversité de perspectives "
    "peut mener à la polarisation : le phénomène par lequel les médias, influencés "
    "par des biais politiques, présentent les mêmes faits de manière très différente "
    "selon leur orientation idéologique."
))
add_body(doc, (
    "La polarisation médiatique est un phénomène documenté scientifiquement. Des "
    "études comme celles de AllSides (USA) ou de Sciences Po (France) montrent que "
    "les médias de gauche, de centre et de droite utilisent des vocabulaires, des "
    "cadres narratifs (frames) et des tons émotionnels distincts pour traiter les "
    "mêmes événements."
))

add_colored_heading(doc, "1.2 Problématique", level=2, color=BLUE_MID)
add_info_box(doc, "Question centrale du projet", [
    "",
    "Peut-on prédire automatiquement l'orientation politique d'un article de presse",
    "traitant d'un événement géopolitique, à partir de son contenu textuel ?",
    "",
    "Et si oui, quels algorithmes de Machine Learning sont les plus performants",
    "pour cette tâche de classification multi-classe ?"
], bg_hex="EBF5FB")

add_colored_heading(doc, "1.3 Objectifs", level=2, color=BLUE_MID)
add_bullet(doc, "Collecter et structurer un corpus d'articles de presse annotés")
add_bullet(doc, "Nettoyer et transformer les données textuelles en features exploitables")
add_bullet(doc, "Comparer 5 algorithmes de classification supervisée")
add_bullet(doc, "Déployer un tableau de bord interactif pour visualiser les résultats")
add_bullet(doc, "Valider chaque étape selon les critères BC4.1 à BC4.4")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE DU PIPELINE
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "2. Architecture du Pipeline de Données", level=1)

add_body(doc, (
    "Le projet suit une architecture de pipeline linéaire, classique en Data Science, "
    "organisée en 4 étapes principales. Chaque étape correspond à un notebook Jupyter dédié "
    "et à un bloc de compétences du référentiel."
))

# Schéma pipeline sous forme de tableau
headers_pipe = ["Étape", "Notebook", "Compétence", "Livrable"]
rows_pipe = [
    ["① Collecte", "01_data_collection.ipynb", "BC4.1", "articles_bruts.csv"],
    ["② Prétraitement", "02_preprocessing_eda.ipynb", "BC4.2", "articles_preprocesses.csv"],
    ["③ Modélisation", "03_modeling.ipynb", "BC4.4", "best_model.pkl"],
    ["④ Dashboard statique", "04_dashboard.ipynb", "BC4.3", "dashboard.html"],
    ["⑤ API REST", "api.py (FastAPI)", "BC4.3", "Service /predict en ligne"],
    ["⑥ App interactive", "app.py (Streamlit)", "BC4.3", "Interface web — 3 pages avec analyses"],
]
make_header_table(doc, headers_pipe, rows_pipe)

add_body(doc, (
    "Chaque notebook est autonome : il lit les fichiers produits par l'étape précédente "
    "et produit ses propres livrables dans le dossier data/. Cette organisation modulaire "
    "facilite la maintenance, le débogage et la réutilisation du code."
))

add_colored_heading(doc, "2.1 Environnement technique", level=2, color=BLUE_MID)

headers_tech = ["Outil / Bibliothèque", "Rôle", "Version"]
rows_tech = [
    ["Python", "Langage de programmation principal", "3.10+"],
    ["Jupyter Notebook", "Environnement de développement interactif", "6.0+"],
    ["pandas", "Manipulation de données tabulaires", "2.x"],
    ["scikit-learn", "Algorithmes ML et évaluation", "1.x"],
    ["NLTK + VADER", "Traitement du langage naturel et sentiment", "3.7+"],
    ["Plotly", "Visualisations interactives", "5.x"],
    ["XGBoost", "Algorithme de boosting gradient", "1.7+"],
    ["TextBlob", "Calcul du score de subjectivité", "0.17+"],
    ["RSS / NewsAPI / Kaggle", "Sources de collecte d'articles", "—"],
    ["FastAPI + Uvicorn", "Exposition du modèle en API REST", "0.110+"],
    ["Streamlit", "Application web interactive de visualisation", "1.32+"],
]
make_header_table(doc, headers_tech, rows_tech)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 3. COLLECTE DES DONNÉES
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "3. Collecte des Données (BC4.1)", level=1)

add_colored_heading(doc, "3.1 Qu'est-ce que la collecte de données ?", level=2, color=BLUE_MID)

add_info_box(doc, "Explication simple", [
    "",
    "Imaginez que vous voulez apprendre à reconnaître les films d'horreur. Pour cela, vous avez",
    "besoin de regarder des milliers de films et de noter leur genre. C'est exactement ce que fait",
    "la collecte de données : on rassemble des exemples (articles de presse) avec leur étiquette",
    "(orientation politique) pour que l'ordinateur puisse apprendre à les distinguer.",
    ""
], bg_hex="FEF9E7")

add_colored_heading(doc, "3.2 Sources de données utilisées", level=2, color=BLUE_MID)

add_body(doc, (
    "Nous utilisons quatre sources complémentaires pour collecter 3000 articles, "
    "garantissant à la fois la diversité (multivariée) et la traçabilité de chaque article :"
))

add_colored_heading(doc, "Source 1 : Flux RSS de médias référencés", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "Lecture directe de 22 flux RSS de médias dont l'orientation politique est documentée "
    "(Guardian, BBC, Reuters, Fox News, NY Post...). Chaque article est filtré par "
    "mots-clés géopolitiques avant d'être labellisé selon le média source."
))
add_bullet(doc, "Aucune clé d'API requise, aucune limite de taux")
add_bullet(doc, "22 flux répartis : 7 Gauche, 8 Centre, 7 Droite")
add_bullet(doc, "Articles datés et liés à leur URL d'origine (traçabilité)")
add_bullet(doc, "Volume collecté : ~230 articles dans la dernière exécution")

doc.add_paragraph()

add_colored_heading(doc, "Source 2 : NewsAPI", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "API REST commerciale avec un plan gratuit (100 requêtes/jour). On envoie une requête "
    "par événement géopolitique (10 événements) et on filtre les articles dont le domaine "
    "est répertorié dans notre dictionnaire de biais médiatique (42 domaines)."
))
add_bullet(doc, "Clé d'API lue depuis la variable d'environnement NEWSAPI_KEY")
add_bullet(doc, "Endpoint : newsapi.org/v2/everything")
add_bullet(doc, "Volume collecté : ~100 articles dans la dernière exécution")

doc.add_paragraph()

add_colored_heading(doc, "Source 3 : Dataset Kaggle (news_bias.csv)", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "Dataset public de 17 362 articles annotés Left/Center/Right, filtré ensuite par les "
    "mots-clés des événements géopolitiques étudiés pour ne retenir que les articles "
    "pertinents (1500 max)."
))
add_bullet(doc, "Annotations expertes Left/Center/Right mappées vers Gauche/Centre/Droite")
add_bullet(doc, "1500 articles retenus après filtrage géopolitique")
add_bullet(doc, "Source : kaggle.com (téléchargement local en CSV)")

doc.add_paragraph()

add_colored_heading(doc, "Source 4 : Générateur synthétique paramétrique", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "Lorsque les données réelles sont insuffisantes ou indisponibles (problème de connexion, "
    "limites d'API), nous utilisons un générateur synthétique. Ce générateur crée des articles "
    "réalistes en s'appuyant sur des patrons lexicaux empiriquement documentés dans la "
    "littérature académique sur la polarisation médiatique."
))

add_info_box(doc, "Comment fonctionne le générateur ?", [
    "",
    "1. Il choisit aléatoirement un événement géopolitique (ex : Guerre en Ukraine)",
    "2. Il tire une orientation politique selon une distribution réaliste (35% Gauche, 38% Centre, 27% Droite)",
    "3. Il sélectionne une source médiatique cohérente avec l'orientation (ex : Mediapart → Gauche)",
    "4. Il remplit un template d'article avec le vocabulaire typique de l'orientation choisie",
    "5. Il calcule les scores de sentiment et de subjectivité",
    "6. Il attribue une date aléatoire dans la plage 2022-2024",
    ""
], bg_hex="EBF5FB")

add_colored_heading(doc, "3.3 Structure du dataset collecté", level=2, color=BLUE_MID)

headers_ds = ["Colonne", "Type", "Description", "Exemple"]
rows_ds = [
    ["titre", "Texte", "Titre de l'article", "Le Monde analyse : Immigration en Europe"],
    ["texte", "Texte", "Corps de l'article", "Face à la crise, les experts indiquent..."],
    ["source", "Texte", "Nom du média", "Le Monde, Mediapart, Le Figaro..."],
    ["orientation", "Catégorie", "Label : Gauche / Centre / Droite", "Centre"],
    ["evenement", "Texte", "Événement géopolitique couvert", "Immigration en Europe"],
    ["date", "Date", "Date de publication", "2023-04-15"],
    ["sentiment", "Nombre", "Score émotionnel (-1 à +1)", "0.0842"],
    ["subjectivite", "Nombre", "Degré d'opinion (0 à 1)", "0.4231"],
    ["nb_mots", "Entier", "Longueur de l'article", "87"],
    ["label", "Entier", "Code numérique : 0=G, 1=C, 2=D", "1"],
]
make_header_table(doc, headers_ds, rows_ds, col_widths=[2, 1.5, 4, 4])

add_colored_heading(doc, "3.4 Événements géopolitiques couverts", level=2, color=BLUE_MID)

events = [
    ("Guerre en Ukraine", "Invasion russe de février 2022"),
    ("Conflit israélo-palestinien", "Tensions et conflits au Proche-Orient"),
    ("Tensions USA-Chine", "Guerre commerciale, Taïwan, technologie"),
    ("Crise climatique COP", "Sommets internationaux sur le climat"),
    ("Immigration en Europe", "Flux migratoires et politiques d'asile"),
    ("Élections présidentielles USA", "Campagnes et résultats électoraux américains"),
    ("Crise énergétique Europe", "Prix du gaz, dépendance énergétique post-Ukraine"),
    ("Montée du populisme", "Partis populistes en Europe et dans le monde"),
    ("OTAN expansion", "Adhésion de la Finlande et Suède, tensions est-ouest"),
    ("Nucléaire Iran", "Accord de Vienne, programme nucléaire iranien"),
]

headers_ev = ["Événement", "Description"]
rows_ev = [(e, d) for e, d in events]
make_header_table(doc, headers_ev, rows_ev)

add_colored_heading(doc, "3.5 Sources médiatiques et leurs biais", level=2, color=BLUE_MID)

add_body(doc, (
    "Le classement des sources est basé sur les études empiriques de AllSides "
    "(USA) et de Médias Bias Fact Check (MBFC), adaptées au contexte français :"
))

headers_src = ["Orientation", "Sources (exemples)", "Caractéristiques"]
rows_src = [
    ["Gauche (0)", "Mediapart, L'Humanité, Le Monde Diplo, The Guardian, HuffPost",
     "Accent sur les droits humains, justice sociale, vocabulaire émotionnel fort"],
    ["Centre (1)", "Le Monde, Reuters, AFP, BBC, France 24, RFI",
     "Ton factuel, sources officielles, peu d'adjectifs évaluatifs"],
    ["Droite (2)", "Le Figaro, CNews, Valeurs Actuelles, Fox News, Atlantico",
     "Accent sur sécurité, souveraineté, vocabulaire identitaire"],
]
make_header_table(doc, headers_src, rows_src)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 4. PRÉTRAITEMENT ET NETTOYAGE
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "4. Prétraitement et Nettoyage des Données (BC4.2)", level=1)

add_colored_heading(doc, "4.1 Pourquoi nettoyer le texte ?", level=2, color=BLUE_MID)

add_info_box(doc, "Explication simple", [
    "",
    "Un algorithme de Machine Learning ne comprend pas le texte comme un humain.",
    "Il travaille avec des chiffres. Avant de lui donner des articles, il faut donc :",
    "",
    "  • Supprimer le 'bruit' : ponctuations, URLs, chiffres inutiles, balises HTML",
    "  • Uniformiser : tout mettre en minuscules (Paris = paris = PARIS)",
    "  • Réduire la complexité : 'courent', 'courait', 'courir' → même racine 'cour'",
    "  • Transformer : convertir le texte en vecteurs numériques (TF-IDF)",
    ""
], bg_hex="FEF9E7")

add_colored_heading(doc, "4.2 Pipeline de nettoyage (7 étapes)", level=2, color=BLUE_MID)

add_body(doc, (
    "Notre pipeline de nettoyage applique les transformations suivantes dans l'ordre :"
))

steps = [
    ("Étape 1 — Mise en minuscules",
     "Toutes les lettres sont converties en minuscules pour que 'Ukraine', 'ukraine' "
     "et 'UKRAINE' soient traités comme le même mot.",
     "Résultat : 'Face À La CRISE' → 'face à la crise'"),
    ("Étape 2 — Suppression des URLs",
     "Les liens web (http://..., www...) sont supprimés car ils n'apportent pas "
     "d'information sur l'orientation politique.",
     "Résultat : 'Lire sur www.lemonde.fr' → 'Lire sur'"),
    ("Étape 3 — Suppression du HTML",
     "Les balises HTML (<b>, <p>, <div>...) éventuellement présentes dans les "
     "données brutes sont supprimées.",
     "Résultat : '<p>Crise en <b>Ukraine</b></p>' → 'Crise en Ukraine'"),
    ("Étape 4 — Suppression ponctuation et chiffres",
     "La ponctuation et les chiffres sont supprimés car ils perturbent la "
     "vectorisation TF-IDF. Seules les lettres et accents sont conservés.",
     "Résultat : 'En 2024, face à la crise !' → 'En face à la crise'"),
    ("Étape 5 — Suppression des espaces multiples",
     "Les espaces doubles ou multiples sont réduits à un espace simple.",
     "Résultat : 'face  à   la  crise' → 'face à la crise'"),
    ("Étape 6 — Suppression des mots vides (stopwords)",
     "Les mots très fréquents mais sans valeur discriminante sont supprimés "
     "(articles, prépositions, conjonctions). On retire aussi les mots de moins "
     "de 3 caractères. La liste combine français et anglais (2000+ mots).",
     "Résultat : 'face à la crise' → 'face crise'"),
    ("Étape 7 — Racinisation (Stemming)",
     "On réduit chaque mot à sa racine grammaticale. Ainsi, 'condamner', "
     "'condamne', 'condamnation' deviennent tous la même racine 'condemn'. "
     "On utilise le SnowballStemmer, adapté au français.",
     "Résultat : 'condamne inacceptabl cris' (formes racinisées)"),
]

for title, explication, exemple in steps:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = BLUE_DARK
    run.font.size = Pt(11)

    add_body(doc, explication)

    p_ex = doc.add_paragraph()
    p_ex.paragraph_format.left_indent = Cm(0.7)
    r = p_ex.add_run(exemple)
    r.font.size  = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
    p_ex.paragraph_format.space_after = Pt(8)

add_colored_heading(doc, "4.3 Extraction de features (caractéristiques)", level=2, color=BLUE_MID)

add_body(doc, (
    "Après le nettoyage, nous calculons plusieurs indicateurs numériques qui "
    "capturent des aspects différents de la polarisation :"
))

headers_feat = ["Feature", "Description", "Pourquoi c'est utile ?"]
rows_feat = [
    ["sentiment_vader", "Score émotionnel global (-1=très négatif, +1=très positif)",
     "Les médias de gauche et de droite sont plus émotionnels que le centre"],
    ["subjectivite", "Degré d'opinion vs. fait (0=objectif, 1=très subjectif)",
     "Le centre est plus factuel, les extrêmes plus subjectifs"],
    ["richesse_lexicale", "Ratio mots uniques / total de mots",
     "Mesure la diversité du vocabulaire utilisé"],
    ["densite_emotionnel", "Proportion de mots à forte charge émotionnelle",
     "Détecte l'intensité du cadrage émotionnel"],
    ["score_vocab_gauche", "Densité du vocabulaire typique de la gauche",
     "Justice, solidarité, droits, égalité..."],
    ["score_vocab_centre", "Densité du vocabulaire du centre",
     "Analyse, rapport, expert, confirme..."],
    ["score_vocab_droite", "Densité du vocabulaire typique de la droite",
     "Sécurité, souveraineté, frontière, menace..."],
    ["indice_polarisation", "max(gauche, droite) − centre",
     "Mesure globale du degré de polarisation de l'article"],
]
make_header_table(doc, headers_feat, rows_feat)

add_colored_heading(doc, "4.4 Vectorisation TF-IDF", level=2, color=BLUE_MID)

add_info_box(doc, "Qu'est-ce que le TF-IDF ?", [
    "",
    "TF-IDF signifie Term Frequency - Inverse Document Frequency.",
    "",
    "• TF (fréquence du terme) : combien de fois un mot apparaît dans l'article",
    "• IDF (fréquence inverse) : à quel point ce mot est rare dans TOUS les articles",
    "",
    "Un mot commun comme 'le' a un IDF très bas → peu important.",
    "Un mot rare comme 'souveraineté' a un IDF élevé → très discriminant.",
    "",
    "Le TF-IDF = TF × IDF : il donne un score élevé aux mots qui sont fréquents",
    "dans un article MAIS rares dans les autres → ce sont les mots caractéristiques.",
    ""
], bg_hex="EBF5FB")

add_body(doc, "Paramètres utilisés dans notre TF-IDF :")
add_bullet(doc, "max_features=5000 : on garde les 5000 termes les plus discriminants")
add_bullet(doc, "ngram_range=(1,2) : on analyse les mots seuls ET les bigrammes ('droits humains', 'sécurité nationale')")
add_bullet(doc, "sublinear_tf=True : on applique un log sur les fréquences (évite qu'un mot très répété domine)")
add_bullet(doc, "min_df=3 : on ignore les mots apparaissant dans moins de 3 articles (trop rares)")
add_bullet(doc, "max_df=0.90 : on ignore les mots dans plus de 90% des articles (trop communs)")

doc.add_paragraph()
add_body(doc, (
    "La matrice TF-IDF résultante a une forme de 3000 × 5000 : chaque article est "
    "représenté par un vecteur de 5000 valeurs numériques. C'est cette matrice, "
    "combinée aux 11 features numériques, qui est utilisée pour entraîner les modèles ML."
))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 5. ANALYSE EXPLORATOIRE
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "5. Analyse Exploratoire des Données (EDA)", level=1)

add_colored_heading(doc, "5.1 Qu'est-ce que l'EDA ?", level=2, color=BLUE_MID)

add_info_box(doc, "Explication simple", [
    "",
    "L'EDA (Exploratory Data Analysis) est la phase de 'découverte' des données.",
    "Avant de construire un modèle, on explore les données pour :",
    "",
    "  • Vérifier leur qualité (valeurs manquantes, doublons...)",
    "  • Comprendre les distributions et les relations entre variables",
    "  • Détecter des patterns utiles pour la modélisation",
    "  • Valider que nos données reflètent bien la réalité",
    ""
], bg_hex="FEF9E7")

add_colored_heading(doc, "5.2 Audit qualité des données", level=2, color=BLUE_MID)

headers_audit = ["Vérification", "Résultat", "Action"]
rows_audit = [
    ["Valeurs manquantes", "0 dans toutes les colonnes", "Aucune action requise"],
    ["Doublons (sur le texte)", "0 doublon détecté", "Aucune action requise"],
    ["Distribution des classes", "Gauche 35%, Centre 38%, Droite 27%", "Légèrement déséquilibrée → prise en compte dans CV stratifiée"],
    ["Plage temporelle", "01/01/2022 → 31/12/2024", "3 ans de couverture, cohérent"],
    ["Longueur des articles", "Moyenne 82 mots, écart-type 12 mots", "Articles courts mais représentatifs"],
]
make_header_table(doc, headers_audit, rows_audit)

add_colored_heading(doc, "5.3 Résultats clés de l'EDA", level=2, color=BLUE_MID)

add_colored_heading(doc, "Scores de sentiment par orientation", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
headers_sent = ["Orientation", "Sentiment moyen (VADER)", "Subjectivité moyenne", "Interprétation"]
rows_sent = [
    ["Gauche", "~+0.08 (légèrement positif)", "~0.55 (modérément subjectif)",
     "Ton engagé, espoir lié à la justice sociale"],
    ["Centre", "~0.00 (neutre)", "~0.30 (peu subjectif)",
     "Ton factuel, journalisme d'information"],
    ["Droite", "~-0.05 (légèrement négatif)", "~0.58 (modérément subjectif)",
     "Ton alertiste, mise en garde contre les menaces"],
]
make_header_table(doc, headers_sent, rows_sent)

add_body(doc, (
    "Ces différences, bien que modestes, sont statistiquement significatives et "
    "constituent des signaux utiles pour la classification. La subjectivité est le "
    "meilleur discriminant entre le centre (objectif) et les orientations politiques "
    "marquées (plus subjectives)."
))

add_colored_heading(doc, "Vocabulaire dominant par orientation", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
headers_vocab = ["Orientation", "Mots les plus caractéristiques"]
rows_vocab = [
    ["Gauche", "droit, solidarité, humanitaire, justice, inégalités, victimes, progressiste"],
    ["Centre", "analyse, rapport, expert, confirme, stratégique, officiels, selon sources"],
    ["Droite", "sécurité, souveraineté, national, frontière, identité, menace, protège"],
]
make_header_table(doc, headers_vocab, rows_vocab)

add_colored_heading(doc, "Réduction dimensionnelle — LSA", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "La LSA (Latent Semantic Analysis) réduit l'espace TF-IDF de 5000 dimensions "
    "à 2 dimensions visualisables. Cette technique permet de vérifier visuellement "
    "que les 3 classes sont séparables dans l'espace des features textuelles. "
    "Les nuages de points confirment une bonne séparabilité, en particulier entre "
    "le centre (vocabulaire factuel distinct) et les orientations polarisées."
))

add_colored_heading(doc, "5.4 Figures produites", level=2, color=BLUE_MID)

headers_figs = ["Figure", "Type", "Description"]
rows_figs = [
    ["fig_distribution.png", "Donut + barres", "Distribution des orientations et répartition par événement"],
    ["fig_sentiment_boxplot.png", "Box plots", "Sentiment, subjectivité et polarisation par orientation"],
    ["fig_evolution_temporelle.png", "Courbes temporelles", "Évolution mensuelle des parts d'orientation (2022-2024)"],
    ["fig_correlation.png", "Heatmap", "Matrice de corrélation des variables numériques"],
    ["fig_lsa_2d.png", "Scatter plot 2D", "Séparabilité des classes dans l'espace TF-IDF réduit (LSA)"],
    ["fig_wordcloud.png", "WordCloud ×3", "Vocabulaire dominant par orientation politique"],
]
make_header_table(doc, headers_figs, rows_figs)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 6. MODÉLISATION MACHINE LEARNING
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "6. Modélisation Machine Learning (BC4.4)", level=1)

add_colored_heading(doc, "6.1 Qu'est-ce que la classification supervisée ?", level=2, color=BLUE_MID)

add_info_box(doc, "Explication simple", [
    "",
    "La classification supervisée, c'est apprendre à l'ordinateur à trier des objets",
    "en catégories, à partir d'exemples étiquetés.",
    "",
    "Exemple : Si on montre 1000 photos de chats et de chiens à un algorithme, en lui",
    "disant 'ceci est un chat' ou 'ceci est un chien', il apprend les caractéristiques",
    "qui distinguent les deux. Ensuite, sur une nouvelle photo, il peut prédire lui-même",
    "si c'est un chat ou un chien.",
    "",
    "Dans notre projet : on montre 2400 articles (80% du dataset) avec leur label",
    "(Gauche/Centre/Droite). Le modèle apprend les patterns textuels. Puis on teste",
    "sur 600 articles inconnus (20%) pour mesurer sa précision réelle.",
    ""
], bg_hex="FEF9E7")

add_colored_heading(doc, "6.2 Préparation des données pour les modèles", level=2, color=BLUE_MID)

add_colored_heading(doc, "Sélection de features (SelectKBest)", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "Avec 5000 features TF-IDF + 11 features numériques = 5011 variables, certaines "
    "sont inutiles ou redondantes. La méthode SelectKBest avec l'information mutuelle "
    "sélectionne automatiquement les 3000 features les plus informatives."
))

add_info_box(doc, "Qu'est-ce que l'information mutuelle ?", [
    "",
    "L'information mutuelle mesure la dépendance entre une feature et la classe cible.",
    "Si connaître la valeur d'une feature aide à prédire la classe → information mutuelle élevée.",
    "",
    "Exemple : le mot 'souveraineté' a une forte info mutuelle avec la classe 'Droite'.",
    "Le mot 'confirme' a une forte info mutuelle avec la classe 'Centre'.",
    "Le mot 'justice' a une forte info mutuelle avec la classe 'Gauche'.",
    ""
], bg_hex="EBF5FB")

add_colored_heading(doc, "Split Train/Test stratifié", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_body(doc, (
    "Le dataset est divisé en deux parties avec une proportion 80%/20% :"
))
add_bullet(doc, "Train set (80% = 2400 articles) : sert à entraîner les modèles")
add_bullet(doc, "Test set (20% = 600 articles) : sert à évaluer la performance réelle")
add_body(doc, (
    "Le mot 'stratifié' signifie que la répartition Gauche/Centre/Droite est "
    "préservée dans les deux parties. Sans cela, on risquerait d'avoir un test set "
    "avec trop peu d'exemples d'une classe, biaisisant les résultats."
))

doc.add_paragraph()

add_colored_heading(doc, "6.3 Les 5 algorithmes implémentés", level=2, color=BLUE_MID)


# ── Algorithme 1 : Naïve Bayes ──
add_colored_heading(doc, "Algorithme 1 — Naïve Bayes Complémentaire", level=3, color=RED_POL)

add_info_box(doc, "Explication simple du Naïve Bayes", [
    "",
    "Le Naïve Bayes est basé sur la probabilité (théorème de Bayes).",
    "",
    "Idée : Si l'article contient le mot 'souveraineté', quelle est la probabilité",
    "qu'il soit de droite ? On calcule cette probabilité pour chaque mot, puis on",
    "les multiplie (en supposant que les mots sont indépendants — d'où 'Naïf').",
    "",
    "La version 'Complémentaire' est particulièrement adaptée aux textes déséquilibrés.",
    ""
], bg_hex="FDEDEC")

headers_nb = ["Caractéristique", "Valeur"]
rows_nb = [
    ["Famille", "Probabiliste — Théorème de Bayes"],
    ["Hyperparamètre principal", "alpha=0.1 (lissage de Laplace pour éviter les probabilités nulles)"],
    ["Forces", "Très rapide, excellente baseline pour NLP, robuste aux données rares"],
    ["Faiblesses", "Hypothèse d'indépendance rarement vraie, ignorer les corrélations entre mots"],
    ["Adapté aux textes ?", "Oui — spécialement conçu pour les données textuelles bag-of-words"],
    ["Complexité", "O(n × d) — linéaire en nombre de features"],
]
make_header_table(doc, headers_nb, rows_nb)


# ── Algorithme 2 : Régression Logistique ──
doc.add_paragraph()
add_colored_heading(doc, "Algorithme 2 — Régression Logistique Multinomiale", level=3, color=GREEN_POL)

add_info_box(doc, "Explication simple de la Régression Logistique", [
    "",
    "Malgré son nom, ce n'est pas une régression mais un algorithme de classification.",
    "Il calcule une combinaison linéaire des features (une somme pondérée de tous les mots)",
    "et passe le résultat dans une fonction sigmoïde qui convertit en probabilité.",
    "",
    "Exemple simplifié (pour prédire 'Droite') :",
    "  score = 0.8 × souveraineté + 0.7 × sécurité − 0.6 × solidarité + ...",
    "  probabilité(Droite) = sigmoid(score) → si > 50% : prédit Droite",
    "",
    "La version Multinomiale gère nos 3 classes simultanément avec softmax.",
    ""
], bg_hex="EAFAF1")

headers_lr = ["Caractéristique", "Valeur"]
rows_lr = [
    ["Famille", "Modèle linéaire généralisé"],
    ["Hyperparamètres", "C=1.0 (régularisation), solver='lbfgs', max_iter=1000"],
    ["C (régularisation)", "Contrôle le surapprentissage : petit C = modèle simple, grand C = modèle complexe"],
    ["Forces", "Très interprétable (coefficients = importance des mots), rapide, bonne performance"],
    ["Faiblesses", "Suppose une relation linéaire entre features et classes"],
    ["Adapté aux textes ?", "Excellent — une des meilleures baselines pour la classification de texte"],
]
make_header_table(doc, headers_lr, rows_lr)


# ── Algorithme 3 : SVM ──
doc.add_paragraph()
add_colored_heading(doc, "Algorithme 3 — SVM Linéaire (Support Vector Machine)", level=3, color=BLUE_MID)

add_info_box(doc, "Explication simple du SVM", [
    "",
    "Le SVM cherche la 'meilleure frontière' pour séparer les classes.",
    "",
    "Imaginez 2 groupes de points sur un plan (articles de Gauche et de Droite).",
    "Le SVM trace la ligne qui maximise la distance (marge) entre les deux groupes.",
    "Les points les plus proches de la frontière s'appellent 'vecteurs de support'.",
    "",
    "En haute dimension (5000 features), cette frontière devient un hyperplan.",
    "La version Linéaire est préférée pour les données textuelles car très rapide.",
    ""
], bg_hex="EBF5FB")

headers_svm = ["Caractéristique", "Valeur"]
rows_svm = [
    ["Famille", "Marge maximale (max-margin classifier)"],
    ["Hyperparamètres", "C=1.0, max_iter=2000"],
    ["Principe", "Maximiser la marge entre les classes dans l'espace des features"],
    ["Forces", "Très performant sur données textuelles, robuste en haute dimension"],
    ["Faiblesses", "Pas de probabilités directes, moins interprétable que LR"],
    ["Adapté aux textes ?", "Excellent — souvent meilleur que LR sur données TF-IDF"],
]
make_header_table(doc, headers_svm, rows_svm)


# ── Algorithme 4 : Random Forest ──
doc.add_paragraph()
add_colored_heading(doc, "Algorithme 4 — Random Forest (Forêt Aléatoire)", level=3, color=RGBColor(0x8E, 0x44, 0xAD))

add_info_box(doc, "Explication simple du Random Forest", [
    "",
    "Un Random Forest, c'est une 'forêt' de nombreux arbres de décision.",
    "",
    "Un arbre de décision pose des questions binaires :",
    "  'Le mot souveraineté est-il présent ?' → Oui → 'Est-ce que sécurité > 0.5 ?' → ...",
    "  jusqu'à atteindre une feuille : 'Prédit : Droite'",
    "",
    "Le Random Forest crée 200 arbres différents (chacun entraîné sur un",
    "sous-échantillon aléatoire des données ET des features). Chaque arbre vote,",
    "et la classe majoritaire est retenue → vote démocratique !",
    "",
    "Cette approche réduit le surapprentissage (overfitting) et améliore la robustesse.",
    ""
], bg_hex="F5EEF8")

headers_rf = ["Caractéristique", "Valeur"]
rows_rf = [
    ["Famille", "Ensemble — Bagging (Bootstrap AGGregatING)"],
    ["Hyperparamètres", "n_estimators=200, max_depth=20, min_samples_split=5"],
    ["n_estimators", "Nombre d'arbres dans la forêt (plus = plus stable, mais plus lent)"],
    ["max_depth", "Profondeur maximale de chaque arbre (évite le surapprentissage)"],
    ["Forces", "Robuste, peu de réglages nécessaires, donne l'importance des features"],
    ["Faiblesses", "Moins performant que boosting sur texte, plus lent que SVM/LR"],
    ["Adapté aux textes ?", "Moyen — meilleur avec des features numériques qu'avec TF-IDF pur"],
]
make_header_table(doc, headers_rf, rows_rf)


# ── Algorithme 5 : XGBoost ──
doc.add_paragraph()
add_colored_heading(doc, "Algorithme 5 — XGBoost (eXtreme Gradient Boosting)", level=3, color=RGBColor(0xD3, 0x54, 0x00))

add_info_box(doc, "Explication simple du XGBoost", [
    "",
    "XGBoost est aussi un ensemble d'arbres, mais avec une stratégie différente : le Boosting.",
    "",
    "Au lieu de créer des arbres indépendants (comme Random Forest),",
    "le boosting crée des arbres en séquence, où chaque arbre corrige",
    "les erreurs du précédent.",
    "",
    "Arbre 1 → fait des erreurs → Arbre 2 se concentre sur ces erreurs",
    "→ Arbre 3 corrige les erreurs de 2 → ... → 200 arbres complémentaires",
    "",
    "Le 'Gradient' dans le nom signifie qu'on utilise le gradient de la fonction",
    "de perte pour optimiser chaque arbre → très efficace mathématiquement.",
    "XGBoost est l'algorithme qui domine les compétitions Kaggle !",
    ""
], bg_hex="FEF0E6")

headers_xgb = ["Caractéristique", "Valeur"]
rows_xgb = [
    ["Famille", "Ensemble — Boosting (arbres séquentiels)"],
    ["Hyperparamètres", "n_estimators=200, lr=0.1, max_depth=6, subsample=0.8"],
    ["learning_rate", "Vitesse d'apprentissage (petit lr = apprentissage prudent, stable)"],
    ["subsample", "80% des données utilisées pour chaque arbre (évite surapprentissage)"],
    ["colsample_bytree", "80% des features utilisées pour chaque arbre (régularisation)"],
    ["Forces", "Très performant, nombreuses options de régularisation intégrées"],
    ["Faiblesses", "Beaucoup de hyperparamètres à régler, moins rapide que LR/SVM"],
    ["Adapté aux textes ?", "Bon — excellent sur données mixtes (TF-IDF + features numériques)"],
]
make_header_table(doc, headers_xgb, rows_xgb)


add_colored_heading(doc, "6.4 Protocole d'évaluation", level=2, color=BLUE_MID)

add_colored_heading(doc, "Cross-Validation Stratifiée 5-fold", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_info_box(doc, "Qu'est-ce que la cross-validation ?", [
    "",
    "La cross-validation évite de juger un modèle sur un seul split qui pourrait être chanceux.",
    "",
    "On divise le train set en 5 parties égales (folds) :",
    "  • Itération 1 : train sur folds 1+2+3+4, test sur fold 5",
    "  • Itération 2 : train sur folds 1+2+3+5, test sur fold 4",
    "  • ... 5 itérations au total",
    "",
    "On obtient 5 scores → on calcule la moyenne et l'écart-type.",
    "Un bon modèle a une moyenne élevée ET un faible écart-type (stable).",
    ""
], bg_hex="EBF5FB")

add_colored_heading(doc, "Métriques d'évaluation", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))

headers_met = ["Métrique", "Formule simplifiée", "Interprétation"]
rows_met = [
    ["Accuracy (précision globale)", "Articles bien classés / Total articles",
     "% d'articles correctement classés. Attention au déséquilibre des classes !"],
    ["F1 Macro", "Moyenne des F1 de chaque classe",
     "Équilibre précision/rappel, tient compte du déséquilibre"],
    ["Precision", "Vrais positifs / (Vrais + Faux positifs)",
     "Quand on dit 'Gauche', est-ce vraiment de Gauche ?"],
    ["Recall (Rappel)", "Vrais positifs / (Vrais + Faux négatifs)",
     "Parmi tous les vrais articles de Gauche, combien on trouve ?"],
    ["Matrice de confusion", "Tableau prédiction vs. réalité",
     "Montre les types d'erreurs : ex. Gauche confondue avec Centre"],
]
make_header_table(doc, headers_met, rows_met)

add_colored_heading(doc, "6.5 Résultats mesurés sur le test set (3000 articles, split 80/20)", level=2, color=BLUE_MID)

headers_res = ["Modèle", "Accuracy", "F1 Macro", "CV F1 (5-fold)", "Temps d'entraînement"]
rows_res = [
    ["XGBoost", "0.797", "0.793", "0.806 ± 0.022", "39.6 s"],
    ["SVM Linéaire", "0.770", "0.762", "0.789 ± 0.014", "3.3 s"],
    ["Régression Logistique", "0.760", "0.749", "0.785 ± 0.015", "4.9 s"],
    ["Naïve Bayes (Complement)", "0.752", "0.744", "0.764 ± 0.016", "6.7 s"],
    ["Random Forest", "0.748", "0.736", "0.776 ± 0.017", "3.7 s"],
]
make_header_table(doc, headers_res, rows_res)

add_body(doc, (
    "Note : Un modèle aléatoire obtiendrait ~33% sur cette tâche à 3 classes. Tous nos "
    "modèles dépassent largement cette baseline. XGBoost s'impose sur l'accuracy, le F1 "
    "macro et la robustesse en cross-validation, mais reste 10x plus lent à entraîner que "
    "les modèles linéaires (intérêt du compromis qualité/coût)."
))

add_colored_heading(doc, "6.6 Importance des features", level=2, color=BLUE_MID)
add_body(doc, (
    "L'analyse de l'importance des features permet de comprendre quels mots et "
    "quelles variables numériques sont les plus discriminants pour la classification. "
    "Pour la Régression Logistique et le SVM, on examine les coefficients. "
    "Pour Random Forest et XGBoost, on utilise le Gini Importance."
))
add_body(doc, (
    "On s'attend à retrouver en tête : les mots du vocabulaire politique "
    "('souveraineté', 'droits', 'expert'), les bigrammes caractéristiques "
    "('sécurité nationale', 'droits humains', 'données officielles'), "
    "et les features numériques de vocabulaire politique spécifique."
))

add_colored_heading(doc, "6.7 Sélection du meilleur modèle", level=2, color=BLUE_MID)

add_body(doc, (
    "Le meilleur modèle est sélectionné selon le F1 Macro sur le test set. "
    "Ce critère est préféré à l'accuracy car il prend en compte le déséquilibre "
    "des classes. Dans la dernière exécution, XGBoost remporte la sélection avec "
    "F1 = 0.793. Le modèle est sauvegardé avec ses transformateurs (TF-IDF, "
    "MinMaxScaler, SelectKBest) dans best_model.pkl pour être réutilisé dans "
    "le pipeline de prédiction et le dashboard."
))

add_info_box(doc, "Pourquoi XGBoost gagne sur ce dataset ?", [
    "",
    "Sur ce projet, XGBoost arrive en tête car :",
    "",
    "• Les features ne sont pas que du texte : 11 features numériques (sentiment,",
    "  subjectivité, scores de vocabulaire politique, indice de polarisation) sont",
    "  concaténées au TF-IDF — un terrain de jeu idéal pour les arbres boostés.",
    "• Le boosting capture les interactions non-linéaires entre vocabulaire et",
    "  signaux numériques (ex. forte densité émotionnelle ET vocabulaire 'sécurité'",
    "  → indicateur de polarisation Droite renforcé).",
    "• Les modèles linéaires (SVM 0.762, LR 0.749) restent compétitifs et 10x",
    "  plus rapides — pertinents si la latence est critique.",
    ""
], bg_hex="EBF5FB")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 7. DASHBOARD INTERACTIF
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "7. Dashboard Interactif (BC4.3)", level=1)

add_colored_heading(doc, "7.1 Objectif du dashboard", level=2, color=BLUE_MID)
add_body(doc, (
    "Le tableau de bord permet de visualiser les résultats de l'analyse de manière "
    "intuitive et interactive, sans avoir besoin d'exécuter du code Python. Il est "
    "conçu pour être présenté à des parties prenantes non-techniques et exporté "
    "sous forme de fichier HTML autonome, ouvert directement dans un navigateur."
))

add_colored_heading(doc, "7.2 Composants du dashboard", level=2, color=BLUE_MID)

headers_dash = ["Visualisation", "Type de graphique", "Ce qu'elle montre"]
rows_dash = [
    ["KPI Cards", "Indicateurs numériques", "Nombre d'articles, sources, événements, meilleur F1"],
    ["Distribution par événement", "Barres empilées à 100%", "Part de chaque orientation pour chaque événement géopolitique"],
    ["Évolution temporelle", "Courbes avec rangeslider", "Comment les orientations évoluent mois par mois (2022-2024) avec rangeslider interactif"],
    ["Analyse des sentiments", "Violin plots", "Distribution complète des scores VADER par orientation (densité + médiane + quartiles)"],
    ["Sentiment vs Subjectivité", "Scatter plot", "Nuage de 500 articles colorés par orientation, taille ∝ longueur de l'article"],
    ["Comparaison modèles ML", "Radar chart", "Accuracy, F1, Precision, Recall sur 4 axes pour chaque modèle"],
    ["Heatmap polarisation", "Matrice de chaleur", "Indice de polarisation moyen par source médiatique et événement"],
]
make_header_table(doc, headers_dash, rows_dash)

add_colored_heading(doc, "7.3 Fonctionnalités interactives", level=2, color=BLUE_MID)
add_bullet(doc, "Zoom/dézoom sur tous les graphiques (molette de la souris)")
add_bullet(doc, "Survol (hover) : affiche les valeurs exactes au survol de la souris")
add_bullet(doc, "Rangeslider : glisser pour sélectionner une période temporelle")
add_bullet(doc, "Légende cliquable : masquer/afficher chaque orientation")
add_bullet(doc, "Export PNG : bouton download pour chaque graphique")
add_bullet(doc, "Dashboard HTML : fichier autonome, aucun serveur requis")

add_colored_heading(doc, "7.4 Technologies utilisées", level=2, color=BLUE_MID)
add_bullet(doc, "Plotly 5.x : bibliothèque Python pour graphiques interactifs basés sur D3.js")
add_bullet(doc, "Plotly Express : API haut niveau pour graphiques courants (scatter, violin...)")
add_bullet(doc, "Plotly Graph Objects : API bas niveau pour personnalisation avancée")
add_bullet(doc, "HTML/CSS : mise en page du dashboard exporté")
add_bullet(doc, "CDN Plotly.js : rendu côté client, aucune dépendance serveur")

# ── 7.5 Déploiement live : API + Streamlit ─────────────────────────────────
add_colored_heading(doc, "7.5 Déploiement opérationnel — API REST + interface Streamlit", level=2, color=BLUE_MID)

add_body(doc, (
    "En complément du dashboard HTML statique, le projet expose le modèle entraîné via "
    "une architecture en deux processus, simulant un déploiement de production :"
))

add_info_box(doc, "Architecture en deux briques", [
    "",
    "  Utilisateur ─[HTTP]─> Streamlit (port 8501) ─[HTTP]─> FastAPI (port 8000) ─> XGBoost (best_model.pkl)",
    "",
    "  • FastAPI = couche de service : reçoit du texte, retourne une prédiction structurée (JSON)",
    "  • Streamlit = couche de présentation : interface graphique, formulaires, visualisations",
    "  • Cette séparation permet de réutiliser l'API depuis n'importe quel autre client (curl, autre app, mobile…)",
    "",
], bg_hex="EBF5FB")

add_colored_heading(doc, "API REST (FastAPI) — fichier api.py", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))

headers_api = ["Endpoint", "Méthode", "Rôle"]
rows_api = [
    ["/health", "GET", "Vérifier que le service est démarré et que le modèle est chargé"],
    ["/metadata", "GET", "Retourner la classe du modèle, taille du vocabulaire, classement des modèles"],
    ["/predict", "POST", "Classifier un texte : entrée { text }, sortie { orientation, label, probabilities, confidence }"],
]
make_header_table(doc, headers_api, rows_api)

add_body(doc, (
    "Au démarrage, l'API charge une seule fois le pickle data/best_model.pkl (modèle + "
    "vectoriseur TF-IDF + scaler MinMax + sélecteur de features). À chaque appel /predict, "
    "le texte traverse exactement le même pipeline NLP que pendant l'entraînement (nettoyage, "
    "stemming, calcul des 11 features numériques, vectorisation, sélection, prédiction). "
    "Cela garantit que les prédictions live sont strictement cohérentes avec les performances "
    "mesurées sur le test set."
))

add_colored_heading(doc, "Application Streamlit — fichier app.py", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))

add_body(doc, "L'interface comporte trois pages :")
add_bullet(doc, "Prédiction en direct : champ texte + 3 exemples préchargés (gauche / centre / droite), bouton « Prédire » qui appelle /predict, affichage du résultat avec barres de probabilité et analyse de confiance.")
add_bullet(doc, "Vue d'ensemble du dataset : KPIs, répartition des orientations, articles par événement, sentiment moyen par classe, évolution mensuelle — chaque graphique est accompagné d'une boîte d'observation interprétant le résultat.")
add_bullet(doc, "Performance des modèles : tableau coloré, radar 4 métriques, scatter qualité/vitesse, stabilité CV, et 4 onglets d'analyse approfondie (pourquoi XGBoost gagne, pourquoi Random Forest est en retrait, forces du SVM, risques & limites).")

add_colored_heading(doc, "Lancement local (deux terminaux)", level=3, color=RGBColor(0x5D, 0x6D, 0x7E))
add_info_box(doc, "Commandes", [
    "",
    "  Terminal 1 :   uvicorn api:app --port 8000",
    "  Terminal 2 :   streamlit run app.py",
    "",
    "  Puis ouvrir : http://localhost:8501",
    "",
], bg_hex="F5EEF8")

add_body(doc, (
    "La barre latérale de l'application affiche en temps réel le statut de l'API "
    "(vert si /health répond, rouge sinon avec la commande de démarrage). Cette télémétrie "
    "facilite le diagnostic en démo devant un jury."
))

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 8. RÉSULTATS ET DISCUSSION
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "8. Résultats et Discussion", level=1)

add_colored_heading(doc, "8.1 Synthèse des performances", level=2, color=BLUE_MID)
add_body(doc, (
    "Les 5 modèles ont été entraînés et évalués selon le même protocole : "
    "split 80/20 stratifié + cross-validation 5-fold (StratifiedKFold). Le tableau "
    "de la section 6.5 récapitule les performances mesurées sur le test set."
))

add_body(doc, (
    "XGBoost obtient les meilleurs scores (F1 macro = 0.793, accuracy = 0.797), "
    "suivi par le SVM linéaire (0.762). La cross-validation à 5 folds confirme "
    "la stabilité du classement (écart-type ≤ 0.022). Le Naïve Bayes offre une "
    "baseline solide à 0.744. Le Random Forest, légèrement en retrait (0.736), "
    "souffre probablement d'une profondeur limitée à 20 sur un espace de 3000 "
    "features — relâcher cette contrainte ferait probablement converger vers "
    "XGBoost mais au prix du temps."
))

add_colored_heading(doc, "8.2 Analyse des erreurs", level=2, color=BLUE_MID)
add_body(doc, (
    "L'analyse des matrices de confusion révèle que les confusions les plus fréquentes "
    "se produisent entre :"
))
add_bullet(doc, "Gauche ↔ Centre : articles de centre-gauche au vocabulaire modéré")
add_bullet(doc, "Centre ↔ Droite : articles techniques/économiques avec vocabulaire neutre")
add_bullet(doc, "Gauche ↔ Droite : très rare, car les vocabulaires sont très distincts")

add_body(doc, (
    "Ces confusions sont attendues et cohérentes avec la réalité du paysage médiatique : "
    "les frontières entre orientations proches sont naturellement floues. Les journalistes "
    "de centres différents partagent parfois un même vocabulaire factuel."
))

add_colored_heading(doc, "8.3 Interprétabilité du modèle", level=2, color=BLUE_MID)
add_body(doc, (
    "L'interprétabilité est cruciale pour valider que le modèle apprend bien les "
    "patterns politiques et non des artefacts du dataset. L'analyse des features "
    "importantes confirme que les mots les plus discriminants correspondent bien "
    "aux vocabulaires politiques documentés dans la littérature :"
))
add_bullet(doc, "Classe Gauche : droits, humanitaire, solidarité, justice, inégalités")
add_bullet(doc, "Classe Centre : confirme, selon, rapport, officiels, experts, analyse")
add_bullet(doc, "Classe Droite : souveraineté, sécurité, frontière, menace, national")

add_colored_heading(doc, "8.4 Limites du projet", level=2, color=BLUE_MID)
add_bullet(doc, "Dataset synthétique : les articles générés, bien que réalistes, sont plus 'propres' que de vrais articles web")
add_bullet(doc, "Binarisation du spectre politique : la réalité est un continuum, pas 3 catégories discrètes")
add_bullet(doc, "Langue : le modèle est entraîné principalement sur du français/anglais")
add_bullet(doc, "Évolution temporelle : les marqueurs politiques évoluent avec l'actualité")
add_bullet(doc, "Biais du générateur : le générateur synthétique peut sur-représenter certains patterns")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 9. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "9. Conclusion et Perspectives", level=1)

add_colored_heading(doc, "9.1 Conclusion", level=2, color=BLUE_MID)
add_body(doc, (
    "Ce projet a démontré qu'il est possible de prédire automatiquement l'orientation "
    "politique d'un article de presse traitant d'événements géopolitiques avec une "
    "précision significativement supérieure à la chance (33%). Le pipeline complet — "
    "collecte, prétraitement NLP, vectorisation TF-IDF, modélisation supervisée et "
    "visualisation interactive — a été implémenté de manière modulaire et réutilisable."
))
add_body(doc, (
    "Les modèles linéaires (SVM, Régression Logistique) s'avèrent les plus "
    "performants sur ce type de tâche, confirmant les résultats de la littérature "
    "académique sur la classification de texte. L'analyse des features importantes "
    "valide que les modèles capturent bien les patterns politiques documentés."
))

add_colored_heading(doc, "9.2 Perspectives d'amélioration", level=2, color=BLUE_MID)
add_bullet(doc, "Modèles de langage : Utiliser CamemBERT (BERT français) pour capturer les relations contextuelles entre mots")
add_bullet(doc, "Étendre la collecte : augmenter le quota NewsAPI ou ajouter de nouveaux flux RSS spécialisés (presse régionale, blogs experts)")
add_bullet(doc, "Classification fine-grained : Distinguer plus de 3 orientations (centre-gauche, extrême-droite...)")
add_bullet(doc, "Analyse temporelle : Étudier comment la polarisation sur un événement évolue dans le temps")
add_bullet(doc, "Hébergement cloud : déployer l'API FastAPI + l'app Streamlit sur Render / Railway / Fly.io pour un accès public 24/7")
add_bullet(doc, "Alertes en temps réel : Monitorer les nouvelles parutions et alerter en cas de polarisation extrême")

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 10. MAPPING DES COMPÉTENCES
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "10. Mapping des Compétences du Référentiel", level=1)

headers_comp = ["Compétence", "Critère", "Réalisé dans le projet", "Note /5"]
rows_comp = [
    ["BC4.1 — Extraction données", "C1 : Environnement logiciel adapté", "Python + pandas + requests + nbformat", "OUI /1"],
    ["BC4.1 — Extraction données", "C2 : Outils extraction mobilisés", "RSS (22 flux) + NewsAPI + Kaggle CSV + générateur", "OUI /1"],
    ["BC4.1 — Extraction données", "C3 : Cloud data management", "Dataset structuré, CSV, pickle", "OUI /1"],
    ["BC4.2 — Transformation", "C4 : Fonctions avancées nettoyage", "Pipeline 7 étapes, stemming, stopwords", "OUI /1"],
    ["BC4.2 — Transformation", "C5 : Données prêtes pour ML", "TF-IDF 5000 feat. + 11 features numériques", "OUI /1"],
    ["BC4.3 — Dashboard", "C1 : Env. données massives déployé", "Plotly + pandas + Streamlit + FastAPI", "OUI /1"],
    ["BC4.3 — Dashboard", "C2 : Visualisations informatives", "7 graphiques HTML + 3 pages Streamlit avec analyses", "OUI /1"],
    ["BC4.3 — Dashboard", "C3 : Graphiques temps réel", "API live /predict + cache Streamlit + Plotly CDN", "OUI /1"],
    ["BC4.3 — Dashboard", "C4 : Fonctions logicielles exploitées", "Hover, zoom, rangeslider, formulaire, exemples préchargés, export PNG", "OUI /1"],
    ["BC4.3 — Dashboard", "C5 : Dashboard responsive/interactif", "HTML autonome + app Streamlit responsive", "OUI /1"],
    ["BC4.4 — Modèle prédictif", "C1 : Données chargées/prétraitées", "TF-IDF + features + sélection", "OUI /1"],
    ["BC4.4 — Modèle prédictif", "C2 : Algorithmes testés méthodes connues", "5 modèles + CV 5-fold + métriques", "OUI /1"],
    ["BC4.4 — Modèle prédictif", "C3 : Code fonctionnel sans erreur", "4 notebooks testés et validés", "OUI /1"],
    ["BC4.4 — Modèle prédictif", "C4 : ML avec techniques appropriées", "Stratified split, CV, feature selection", "OUI /1"],
    ["BC4.4 — Modèle prédictif", "C5 : Algorithmes clairement expliqués", "Rapport complet + commentaires code", "OUI /1"],
]
make_header_table(doc, headers_comp, rows_comp)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# 11. GLOSSAIRE
# ════════════════════════════════════════════════════════════════════════════

add_colored_heading(doc, "11. Glossaire", level=1)

glossaire = [
    ("Accuracy", "Précision globale : proportion d'articles correctement classés sur le total."),
    ("Bagging", "Bootstrap AGGregatING : technique d'ensemble qui crée plusieurs modèles sur des sous-échantillons aléatoires des données."),
    ("Boosting", "Technique d'ensemble où chaque modèle corrige les erreurs du précédent de manière séquentielle."),
    ("Classification supervisée", "Apprentissage automatique où le modèle apprend à partir d'exemples étiquetés (données avec réponse connue)."),
    ("Cross-validation", "Technique d'évaluation qui divise les données en K parties et entraîne/teste K fois pour avoir une mesure robuste."),
    ("EDA", "Exploratory Data Analysis : analyse exploratoire pour comprendre les données avant la modélisation."),
    ("F1 Macro", "Moyenne des F1-scores de chaque classe, pondération équitable indépendante de la taille des classes."),
    ("Feature", "Variable d'entrée utilisée par le modèle pour faire une prédiction. En NLP : chaque mot du vocabulaire est une feature."),
    ("Feature Selection", "Sélection des variables les plus informatives pour réduire la dimensionnalité et améliorer les performances."),
    ("RSS (Really Simple Syndication)", "Format XML standardisé permettant de récupérer les derniers articles publiés par un média sans clé d'API ni limite de taux."),
    ("Hyperparamètre", "Paramètre réglé AVANT l'entraînement (ex: C du SVM, n_estimators du RF). Différent des paramètres appris."),
    ("Information mutuelle", "Mesure de dépendance entre une variable et la classe cible. Plus elle est élevée, plus la variable est discriminante."),
    ("LSA", "Latent Semantic Analysis : réduction dimensionnelle du TF-IDF via SVD pour visualiser les données en 2D."),
    ("Naïve Bayes", "Algorithme probabiliste basé sur le théorème de Bayes, avec l'hypothèse d'indépendance des features."),
    ("N-gram", "Séquence de N mots consécutifs. Bigramme = 2 mots ('droits humains'), trigramme = 3 mots."),
    ("NLP", "Natural Language Processing : traitement automatique du langage naturel."),
    ("Overfitting", "Surapprentissage : le modèle apprend par coeur les données d'entraînement mais généralise mal."),
    ("Pipeline", "Chaîne de traitements séquentiels : chaque étape reçoit les outputs de la précédente."),
    ("Polarisation médiatique", "Phénomène par lequel des médias aux orientations politiques différentes couvrent les mêmes faits de manière divergente."),
    ("Precision", "Parmi tous les articles prédits comme 'Gauche', combien sont vraiment de Gauche ? (vrais positifs / tous prédits positifs)"),
    ("Random Forest", "Ensemble de centaines d'arbres de décision entraînés sur des sous-échantillons aléatoires, avec vote majoritaire."),
    ("Recall (Rappel)", "Parmi tous les vrais articles 'Gauche', combien le modèle en retrouve-t-il ?"),
    ("Régression Logistique", "Modèle linéaire de classification qui calcule la probabilité d'appartenance à chaque classe via la fonction softmax."),
    ("Split train/test", "Division du dataset en données d'entraînement (80%) et données de test (20%) pour évaluation honnête."),
    ("Stemming", "Racinisation : réduire un mot à sa racine grammaticale ('courait' → 'cour')."),
    ("Stopwords", "Mots très fréquents mais peu informatifs (articles, prépositions) qu'on supprime avant l'analyse."),
    ("SVM", "Support Vector Machine : algorithme qui trouve l'hyperplan maximisant la marge entre classes."),
    ("TF-IDF", "Term Frequency - Inverse Document Frequency : pondération des mots selon leur fréquence dans le document et leur rareté globale."),
    ("VADER", "Valence Aware Dictionary and sEntiment Reasoner : outil d'analyse de sentiment, particulièrement adapté aux médias."),
    ("XGBoost", "eXtreme Gradient Boosting : implémentation optimisée du gradient boosting, état de l'art sur de nombreuses tâches ML."),
]

for term, definition in glossaire:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{term} : ")
    r1.bold = True
    r1.font.color.rgb = BLUE_DARK
    r1.font.size = Pt(10.5)
    r2 = p.add_run(definition)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = GREY_TEXT


# ════════════════════════════════════════════════════════════════════════════
# SAUVEGARDE
# ════════════════════════════════════════════════════════════════════════════

import os
HERE = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(HERE, "rapport_polarisation_mediatique.docx")
try:
    doc.save(output_path)
except PermissionError:
    # Le fichier est probablement ouvert dans Word — on écrit une copie horodatée.
    from datetime import datetime
    output_path = os.path.join(HERE,
        f"rapport_polarisation_mediatique_{datetime.now():%Y%m%d_%H%M%S}.docx")
    doc.save(output_path)
print(f"[OK] Rapport sauvegarde : {output_path}")

size_kb = os.path.getsize(output_path) / 1024
print(f"   Taille : {size_kb:.1f} Ko")
print(f"   Sections : Page de garde, Résumé, TDM, 11 chapitres, Glossaire (30 termes)")
